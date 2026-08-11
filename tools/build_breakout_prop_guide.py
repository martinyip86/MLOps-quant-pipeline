from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "deliverables" / "Breakout_Prop_策略基礎學習手冊.docx"

FONT_LATIN = "Arial Unicode MS"
FONT_CJK = "Arial Unicode MS"
INK = "1D2939"
BLUE = "175CD3"
BLUE_DARK = "1849A9"
BLUE_PALE = "EFF8FF"
GRAY = "475467"
GRAY_PALE = "F2F4F7"
GRAY_LINE = "D0D5DD"
GREEN = "027A48"
GREEN_PALE = "ECFDF3"
AMBER = "B54708"
AMBER_PALE = "FFFAEB"
RED = "B42318"
RED_PALE = "FEF3F2"
WHITE = "FFFFFF"


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def apply_font(run, *, size=None, color=INK, bold=None, italic=None, cjk=FONT_CJK):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_para_format(p, before=0, after=6, line=1.25, keep_next=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_next


def add_text(doc, text="", *, size=11, color=INK, bold=False, italic=False,
             before=0, after=6, line=1.25, align=None, keep_next=False):
    p = doc.add_paragraph()
    set_para_format(p, before, after, line, keep_next)
    if align is not None:
        p.alignment = align
    if text:
        apply_font(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_para(doc, chunks, *, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    set_para_format(p, before, after, line)
    if align is not None:
        p.alignment = align
    for text, kwargs in chunks:
        apply_font(p.add_run(text), size=kwargs.get("size", 11),
                   color=kwargs.get("color", INK), bold=kwargs.get("bold"),
                   italic=kwargs.get("italic"))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_numbering(doc: Document):
    numbering = doc.part.numbering_part.element

    def add_abs(abs_id: int, num_fmt: str, text: str, left: int, hanging: int):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        nsid = OxmlElement("w:nsid")
        nsid.set(qn("w:val"), f"{abs_id:08X}")
        abstract.append(nsid)
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        ppr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        lvl.append(ppr)
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), FONT_LATIN)
        rfonts.set(qn("w:hAnsi"), FONT_LATIN)
        rfonts.set(qn("w:eastAsia"), FONT_CJK)
        rpr.append(rfonts)
        lvl.append(rpr)
        abstract.append(lvl)
        numbering.append(abstract)

    def add_num(num_id: int, abs_id: int):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abs_id))
        num.append(abstract_num_id)
        numbering.append(num)

    add_abs(71, "bullet", "•", 540, 270)
    add_num(71, 71)
    add_abs(72, "decimal", "%1.", 540, 270)
    add_num(72, 72)


def add_list_item(doc, text, *, numbered=False, bold_lead=None):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "72" if numbered else "71")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
    set_para_format(p, 0, 4, 1.25)
    if bold_lead and text.startswith(bold_lead):
        apply_font(p.add_run(bold_lead), size=11, color=INK, bold=True)
        apply_font(p.add_run(text[len(bold_lead):]), size=11, color=INK)
    else:
        apply_font(p.add_run(text), size=11, color=INK)
    return p


def add_callout(doc, label, body, *, tone="blue"):
    palette = {
        "blue": (BLUE_PALE, BLUE_DARK),
        "amber": (AMBER_PALE, AMBER),
        "red": (RED_PALE, RED),
        "green": (GREEN_PALE, GREEN),
        "gray": (GRAY_PALE, GRAY),
    }
    fill, accent = palette[tone]
    p = doc.add_paragraph()
    set_para_format(p, 5, 8, 1.20)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "180")
    ind.set(qn("w:right"), "120")
    p_pr.append(ind)
    apply_font(p.add_run(f"{label}　"), size=10.5, color=accent, bold=True)
    apply_font(p.add_run(body), size=10.5, color=INK)
    return p


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    apply_font(run, size=9, color=GRAY)


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_format(hp, 0, 0, 1.0)
    apply_font(hp.add_run("BREAKOUT PROP｜策略基礎學習手冊"), size=9, color=GRAY, bold=True)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), GRAY_LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_table(doc, headers, rows, widths, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, label in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, "E8EEF5")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p, 0, 0, 1.10)
        apply_font(p.add_run(label), size=9.5, color=BLUE_DARK, bold=True)
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=100, bottom=100)
            if r_idx % 2 == 1:
                shade_cell(cell, "F9FAFB")
            p = cell.paragraphs[0]
            p.alignment = (alignments[i] if alignments else
                           (WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER))
            set_para_format(p, 0, 0, 1.10)
            apply_font(p.add_run(str(value)), size=9.5, color=INK)
    set_table_geometry(table, widths)
    add_text(doc, "", after=4)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_numbering(doc)

    # Cover: editorial-cover pattern, compact-reference-guide preset.
    add_text(doc, "學習手冊｜2026 年 7 月版", size=10, color=AMBER, bold=True,
             before=78, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Breakout Prop", size=29, color=BLUE_DARK, bold=True,
             after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "制定交易策略的基礎知識", size=20, color=INK, bold=True,
             after=14, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "從規則理解、風險預算、策略規格化，到回測與執行紀律",
             size=12, color=GRAY, after=62, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(
        doc,
        "核心觀念",
        "Prop 考核首先是「風險預算問題」，其次才是獲利速度問題。沒有時間限制時，最有價值的能力不是加快交易，而是拒絕沒有優勢的交易。",
        tone="blue",
    )
    add_text(doc, "適用對象：準備參加 Breakout Evaluation、但尚未完成可量化策略的新手至進階交易者",
             size=9.5, color=GRAY, before=30, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "資料核對日期：2026-07-27｜規則以你的帳戶 Dashboard、最新官方條款及合約為準",
             size=9, color=GRAY, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    add_heading(doc, "先讀這一頁：用途、界線與完成標準", 1)
    add_text(doc, "這不是報單、投資建議或保證通過的捷徑。它是一個學習框架，幫你把自己的市場觀察轉成可測試、可重複、能在 Breakout 風險限制內生存的策略。")
    add_callout(
        doc,
        "合規紅線",
        "Breakout 的官方規則禁止使用第三方／現成／被行銷為可通過考核的策略，也禁止複製社群、分析師、研究報告或訊號來源的交易想法。請只用本手冊學習方法；你的進出場邏輯必須由你自行提出、驗證、保存版本紀錄，並與實際交易一致。",
        tone="red",
    )
    add_heading(doc, "你完成本手冊後，應能回答 7 個問題", 2)
    competencies = [
        "我的策略在哪一種市場狀態下才有優勢？",
        "進場、失效、停損、出場與不交易條件，能否讓另一個人無歧義地重現？",
        "計入手續費、滑價與持倉費後，期望值是否仍為正？",
        "最差連敗與最大回撤，是否能被所選方案的風險預算容納？",
        "同時持有高度相關部位時，總風險是多少？",
        "在 00:30 UTC 每日重算前後，我的帳戶 equity 最壞情境是多少？",
        "我的交易紀錄能否證明：考核與通過後使用的是同一套原創策略？",
    ]
    for item in competencies:
        add_list_item(doc, item)
    add_heading(doc, "建議學習順序", 2)
    for item in [
        "先理解官方規則與帳戶的真正「死亡線」。",
        "再學期望值、R 倍數、倉位與回撤。",
        "把市場觀察寫成策略規格，而不是憑感覺下單。",
        "用歷史樣本與模擬執行驗證；最後才考慮購買考核。",
    ]:
        add_list_item(doc, item, numbered=True)

    page_break(doc)
    add_heading(doc, "1｜Breakout 規則：先知道遊戲怎麼結束", 1)
    add_text(doc, "以下為官方頁面於 2026-07-27 可見的 1-Step 方案摘要。數字不是你的風控目標，而是違約邊界；實際交易必須保留緩衝。")
    add_table(
        doc,
        ["方案", "獲利目標", "每日最大損失", "最大回撤", "目標／回撤"],
        [
            ["Classic", "10%", "3%", "6% static", "1.67"],
            ["Pro", "12%", "3%", "5% static", "2.40"],
            ["Turbo", "9%", "3%", "3% static", "3.00"],
        ],
        [1800, 1800, 2100, 1900, 1760],
    )
    add_text(doc, "表格解讀：目標／回撤比越高，代表每允許損失 1%，要賺到的百分比越多；這不是報酬風險比，而是考核難度的一個粗略視角。", size=9.5, color=GRAY, before=0, after=8)
    add_heading(doc, "兩條同時生效的 equity 限制", 2)
    add_rich_para(doc, [
        ("每日最大損失：", {"bold": True, "color": BLUE_DARK}),
        ("在每日 00:30 UTC，以當時「不含未平倉部位」的 balance 為基準，1-Step 的 equity 下限為 balance × 97%。未實現損益會計入 equity；只要碰到或越過下限，即屬 breach。", {}),
    ])
    add_rich_para(doc, [
        ("最大回撤：", {"bold": True, "color": BLUE_DARK}),
        ("1-Step 為 static，以初始帳戶為基準：Classic 6%、Pro 5%、Turbo 3%。這條線不會因獲利上移。", {}),
    ])
    add_callout(
        doc,
        "例：100K Classic",
        "初始最大回撤線固定為 94,000。若 00:30 UTC 的 balance 是 105,000，接下來 24 小時的每日 equity 下限為 101,850。此時真正有效的下限是兩者中較高的一條，也就是 101,850。",
        tone="amber",
    )
    add_heading(doc, "成本、槓桿與執行限制", 2)
    for item in [
        "交易費：每邊 0.04% 的名目部位，完整開平約 0.08%，尚未包含滑價。",
        "持倉費：每日合計 0.033%／每個未平倉部位；Breakout Terminal 分成每 4 小時一次、一天 6 次收取。",
        "槓桿：BTC、ETH 最高 5x；其他標的 2x，自動套用。",
        "所有標的 24/7（維護或中斷除外）；允許新聞交易與週末持倉，但流動性與滑價風險仍存在。",
        "平台不支援 partial fills；單一部位無法直接設定分批 take-profit，可用多個部位或手動減倉模擬。",
        "新購帳戶已轉向 Breakout Terminal；舊 DXtrade 帳戶依官方公告維持原狀。",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "2｜策略的六個地基", 1)
    add_text(doc, "策略不是「看到某個形狀就買」。完整策略至少要同時定義下列六項；缺一項，回測與實盤就會逐漸變成兩套東西。")
    foundations = [
        ("市場假設", "你認為哪一種可重複行為造成可交易優勢？必須可被反證。"),
        ("市場狀態", "趨勢、盤整、高／低波動、流動性時段。策略只在適合的狀態啟用。"),
        ("進場觸發", "可觀察、明確、在收盤或成交時能判定的條件；避免「感覺很強」。"),
        ("失效與停損", "市場出現甚麼證據代表原假設錯了？停損應由結構決定，倉位再配合停損。"),
        ("獲利與管理", "固定目標、追蹤、時間出場或條件出場；預先寫明，不因盈虧臨時變更。"),
        ("風險與暫停", "每筆、每日、相關性群組、連敗後降風險，以及何時停止策略。"),
    ]
    add_table(doc, ["地基", "你必須寫清楚的內容"], foundations, [2100, 7260],
              [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    add_heading(doc, "從「想法」到「規則」的判斷測試", 2)
    tests = [
        "二元性：每個條件在當下只能是符合／不符合，不應依賴事後看圖。",
        "可重現性：隔一週重看同一段資料，判定結果應相同。",
        "可執行性：策略要使用 Breakout Terminal 真正可取得的價格與訂單功能。",
        "成本後成立：加入交易費、持倉費與保守滑價後，仍有合理正期望值。",
        "容量合理：名目倉位與流動性不能讓你的成交假設不切實際。",
    ]
    for item in tests:
        add_list_item(doc, item)
    add_callout(doc, "好規則的語氣", "使用「若 A、B、C 同時成立，則在 X 時點採取 Y；若 D 出現則不交易」。避免「通常、應該、看起來、差不多」等無法回測的詞。", tone="green")

    add_heading(doc, "市場狀態比指標名稱更重要", 2)
    add_text(doc, "同一個訊號在趨勢與盤整中可能得到相反結果。先用簡單、可量化的狀態分類，例如：方向（價格相對長期基準）、波動（近期真實波幅的分位數）、流動性（時段與價差）。不要一次堆疊十個高度相關的指標。")

    add_heading(doc, "3｜風險數學：先讓策略有活下去的機會", 1)
    add_heading(doc, "R 倍數與期望值", 2)
    add_rich_para(doc, [
        ("1R", {"bold": True, "color": BLUE_DARK}),
        (" 是你在進場前願意承擔的單筆計畫損失。若一筆虧損剛好等於計畫損失，結果是 −1R；獲利是計畫損失的兩倍，結果是 +2R。", {}),
    ])
    add_callout(doc, "期望值公式", "E(R) = 勝率 × 平均獲利 R − 敗率 × 平均虧損 R − 平均成本 R。只有樣本外、成本後的 E(R) 才有決策價值。", tone="blue")
    add_text(doc, "例：勝率 45%、平均獲利 1.6R、平均虧損 1R、平均成本 0.08R，則 E = 0.45×1.6 − 0.55×1 − 0.08 = 0.09R／筆。正值不代表每月獲利，仍會有長連敗與回撤。")
    add_heading(doc, "倉位計算", 2)
    add_callout(doc, "名目倉位", "可承擔美元風險 ÷（停損距離% ＋ 預估完整交易成本%）＝ 名目部位。滑價與持倉費要依你的持有時間另加緩衝。", tone="gray")
    add_text(doc, "假設 100K 帳戶每筆自訂風險 0.25%（$250），停損距離 1.0%，估計完整成本 0.12%，則名目部位約為 $250 ÷ 1.12% = $22,321。這只是計畫值，實際成交仍可能因跳價或流動性超出預算。")
    add_heading(doc, "把官方上限轉成你的內部風控", 2)
    add_table(
        doc,
        ["層級", "官方邊界", "保守起始值（示例）", "目的"],
        [
            ["單筆風險", "未規定固定值", "0.25% 初始 balance", "容納正常連敗"],
            ["單日自停", "3% breach", "1.0%～1.5%", "保留錯單／滑價空間"],
            ["同時開放風險", "受 equity 限制", "0.50%～0.75%", "控制相關部位"],
            ["降風險門檻", "無", "從峰值 −2R", "避免情緒放大損失"],
            ["策略暫停", "無", "觸及統計異常或流程偏差", "先診斷再恢復"],
        ],
        [1550, 1950, 2300, 3560],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_text(doc, "上表是教學起點，不是個人化建議或通過保證。你的數字應由最差歷史區段、成本、交易頻率與心理承受力共同決定。", size=9.5, color=GRAY)
    add_heading(doc, "連敗預算", 2)
    add_text(doc, "若每筆風險為 0.25%，單純計算下 3%／5%／6% 分別可容納 12／20／24 個完整 −1R。實際可容納數更少，因為手續費、滑價、持倉費、相關部位與超出停損成交都會消耗預算。實務上應讓「最差合理連敗」最多只吃掉約一半最大回撤，另一半留給模型誤差與執行風險。")

    add_heading(doc, "4｜Prop 專用風控：避免策略沒錯、帳戶先死", 1)
    add_heading(doc, "每日重算的操作流程", 2)
    for item in [
        "在 00:30 UTC 前，記錄目前 balance、equity、未實現損益及所有停損後最壞損失。",
        "重算後，以 Dashboard 顯示的每日 equity 下限為唯一權威數字。",
        "計算「今日安全距離」＝目前 equity − 較高的有效下限；再扣除未成交訂單與開放部位最壞損失。",
        "若剩餘距離小於一筆標準風險加成本緩衝，今天停止新增部位。",
    ]:
        add_list_item(doc, item, numbered=True)
    add_callout(doc, "午夜風險", "不要假設昨天的獲利會自動讓今天更安全。每日限制依 00:30 UTC 的 balance 重算，而 open P&L 會在整個 24 小時內影響 equity。跨重算點持倉時，必須重新檢查最壞情境。", tone="amber")
    add_heading(doc, "相關性：三個代幣不等於三個獨立交易", 2)
    add_text(doc, "BTC、ETH、SOL 同方向部位在市場急跌時可能一起受損。把高度相關的部位當成一個「風險群組」，群組總停損風險不得超過你的同時開放風險。若共同因子相同，不要用標的數量製造分散錯覺。")
    add_heading(doc, "考核與 funded 必須是同一個人格", 2)
    for item in [
        "同一套市場、時間框架、進出場邏輯與風險框架。",
        "允許的變更只能是預先定義的版本更新，並保留原因、日期與重新測試結果。",
        "不要在考核用高風險「搏一次」，通過後才換成另一套慢策略；官方條款明文禁止。",
        "不要複製他人訊號、研究報告、社群觀點或現成策略；也不要跨帳戶做相反方向對沖。",
    ]:
        add_list_item(doc, item)
    add_heading(doc, "策略之外的 breach 風險清單", 2)
    for item in [
        "利用報價錯誤、延遲、非公開資訊或前置交易。",
        "帳戶分享、共用憑證，以及同一家庭／裝置／IP 的多帳戶交易風險。",
        "跨帳戶或跨交易者對沖；單一帳戶內的 hedge 也有嚴格條件與公司裁量。",
        "超過平台顯示的單一標的最大部位。",
        "VPN 用於隱瞞或虛報司法管轄地。",
    ]:
        add_list_item(doc, item)
    add_callout(doc, "不確定就先問", "若你的自動化、同住家人、多帳戶、對沖、資料來源或策略靈感可能落入灰區，先以書面向 Breakout Support 說明具體流程並保存回覆，再決定是否交易。", tone="red")

    add_heading(doc, "5｜把原創想法寫成可測試的策略規格", 1)
    add_text(doc, "以下是空白規格，不提供任何第三方進出場邏輯。請用你自己的觀察填寫；每個欄位都應在開始回測前固定。")
    fields = [
        ("策略名稱／版本", "例如：自訂名稱 v0.1；每次變更建立新版本。"),
        ("市場假設", "可被反證的一句話；說明為何可能存在優勢。"),
        ("標的與時間框架", "只列你實際會交易與能取得可靠資料的市場。"),
        ("啟用市場狀態", "趨勢、波動、流動性與時段的量化條件。"),
        ("進場觸發", "所有必要條件、判定時點、訂單類型與有效期限。"),
        ("不交易條件", "價差、低流動性、事件、資料缺漏、當日風控等。"),
        ("初始停損／失效", "價格或狀態條件；不得以『不想虧』取代。"),
        ("出場／管理", "獲利、追蹤、時間、反向訊號；是否允許手動減倉。"),
        ("風險預算", "單筆、每日、同時開放、相關群組、連敗後調整。"),
        ("成本模型", "開平費、滑價、持倉費、未成交與跳價假設。"),
        ("暫停與退役", "何種統計或流程偏差會觸發審查／停止。"),
    ]
    add_table(doc, ["欄位", "你的規格必須包含"], fields, [2500, 6860],
              [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    add_heading(doc, "變更控制", 2)
    add_text(doc, "回測中不要一看到虧損就改規則。先完成預定樣本，再把問題分類為：資料錯誤、規則歧義、成本低估、特定 regime 失效，或純粹隨機波動。任何修正都建立新版本，重新做樣本外測試。")

    add_heading(doc, "6｜驗證流程：證明不是運氣或過度擬合", 1)
    add_heading(doc, "第一階段：資料與回測", 2)
    for item in [
        "資料品質：時間戳、缺值、成交價與 Breakout 原生價格差異；避免用未來資料。",
        "預先定義：先寫規則、成本、樣本區間與主要評估指標，再看結果。",
        "樣本切分：開發區間、驗證區間、完全不碰的樣本外區間。",
        "Regime 覆蓋：上漲、下跌、盤整、高低波動與不同流動性時段。",
        "敏感度：參數稍微改動仍應合理；只在單一精確數值有效是警訊。",
        "保守成交：限價未成交、滑價、費用、持倉費與延遲都要建模。",
    ]:
        add_list_item(doc, item)
    add_heading(doc, "第二階段：你至少要看的指標", 2)
    add_table(
        doc,
        ["指標", "用途", "警訊"],
        [
            ["交易次數", "判斷統計可信度", "樣本很少卻下強結論"],
            ["成本後期望值", "每筆平均優勢", "只看毛利、不含費用"],
            ["Profit factor", "總獲利／總虧損", "少數極端贏家支撐全部績效"],
            ["最大回撤", "估計資金壓力", "接近方案 breach 邊界"],
            ["最長連敗", "設計風險與心理預案", "回測未涵蓋壞 regime"],
            ["MAE／MFE", "檢查停損與管理", "出場依賴事後最佳化"],
            ["月／regime 分布", "檢查穩定性", "績效集中於單一月份"],
        ],
        [1900, 3300, 4160],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_heading(doc, "第三階段：Monte Carlo 與壓力測試", 2)
    add_text(doc, "把歷史交易結果重新排列或依保守分布模擬多條 equity path，觀察在不同交易順序下的 breach 機率。至少把滑價加倍、費用上升、勝率下降、平均獲利縮小、連續未成交與相關部位同時停損納入壓力測試。")
    add_callout(doc, "通過門檻的思路", "不要問「這套回測能不能達到 10%？」；要問「在保守成本與不友善交易順序下，有多少路徑會先碰到 3%／5%／6% 的限制？」", tone="blue")

    add_heading(doc, "7｜模擬執行與交易日誌", 1)
    add_text(doc, "歷史回測驗證邏輯；模擬／forward test 驗證你能否在當下、使用真實操作流程做到。建議至少跨越多種市場狀態，而不是只等固定天數。")
    add_heading(doc, "每筆交易必填", 2)
    for item in [
        "日期時間（UTC 與本地）、標的、方向、策略版本。",
        "市場狀態標籤、符合的進場條件、不交易條件檢查。",
        "進場、停損、預定出場、名目部位、計畫風險 R。",
        "實際費用、持倉費、滑價、成交差異。",
        "結果（美元、%、R）、MAE、MFE、持有時間。",
        "是否完全依規則；若否，偏差類型與立即改正措施。",
        "決策截圖與當時文字理由；避免只保存事後漂亮圖。",
    ]:
        add_list_item(doc, item)
    add_heading(doc, "每日收盤檢查", 2)
    checklist = [
        "□ Dashboard 的每日與總回撤線已記錄",
        "□ 所有 open risk 與相關性群組已重算",
        "□ 今日是否觸及自訂停止線",
        "□ 交易偏差是否已標記，而不是藏在備註裡",
        "□ 下一次 00:30 UTC 重算前的跨夜風險已處理",
        "□ 沒有使用外部訊號、複製想法或臨時改策略",
    ]
    for item in checklist:
        add_text(doc, item, size=10.5, before=0, after=4)
    add_heading(doc, "每 20 筆的檢討節奏", 2)
    add_text(doc, "分開評估「策略品質」與「執行品質」。策略品質看成本後期望值、回撤、regime 與敏感度；執行品質看漏單、追價、移動停損、過量交易與風控違規。不要用執行錯誤推翻策略，也不要用策略正常連敗替執行錯誤找藉口。")

    add_heading(doc, "8｜30 天學習與準備路線圖", 1)
    add_text(doc, "這 30 天不是保證可以購買考核，而是用來建立最低限度的研究與執行能力。若證據不足，就延長，不要為了日期硬上。")
    add_table(
        doc,
        ["階段", "核心任務", "交付物／過關條件"],
        [
            ["第 1–3 天", "逐條閱讀官方 Program Rules、Evaluation Agreement、Dashboard；換算你的 UTC 重算時間。", "一頁規則摘要；能手算每日與總 equity 下限。"],
            ["第 4–7 天", "學 R、期望值、倉位、成本、回撤、相關性。", "自製風險計算表；能解釋最壞情境。"],
            ["第 8–12 天", "只從自己的市場觀察提出假設；寫 v0.1 規格。", "完整填妥策略規格；零模糊詞。"],
            ["第 13–20 天", "回測、切分樣本、加入 Breakout 成本與執行限制。", "成本後報告；按 regime 分解。"],
            ["第 21–24 天", "樣本外、敏感度、Monte Carlo、壓力測試。", "保守情境下 breach 風險可接受。"],
            ["第 25–30 天", "模擬執行、日誌、00:30 UTC 風控演練。", "連續依規則執行；無重大偏差。"],
        ],
        [1500, 4050, 3810],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_heading(doc, "購買考核前的 Go／No-Go", 2)
    go_items = [
        "我已完成 KYC 可行性與所在地資格檢查。",
        "我能在不看答案的情況下計算兩條 breach 線。",
        "策略由我自行提出，沒有複製第三方訊號或現成規則。",
        "成本後、樣本外期望值為正，且不是由極少數交易支撐。",
        "最差 20 筆區段與壓力測試能容納於方案風險預算內，仍有緩衝。",
        "我已用模擬環境證明能遵守單筆、單日、相關性與停止線。",
        "考核費完全可承受歸零，失敗後不會立刻報復性重買。",
    ]
    for item in go_items:
        add_text(doc, f"□ {item}", size=10.5, before=0, after=4)
    add_callout(doc, "No-Go 原則", "任何一項答「否」，就不是提高槓桿或換便宜方案，而是回到相應階段補證據。", tone="red")

    add_heading(doc, "附錄 A｜一頁式交易前卡", 1)
    add_heading(doc, "開盤／開倉前", 2)
    pretrade = [
        "Dashboard 今日 daily equity limit：__________",
        "Static max drawdown limit：__________",
        "目前 equity：__________　今日安全距離：__________",
        "策略名稱／版本：__________　市場狀態：__________",
        "本筆計畫風險：_____%／$__________　相關群組總風險：__________",
        "進場條件全部成立：□　不交易條件全部排除：□",
        "停損／失效：__________　出場規則：__________",
        "含費用與滑價的名目部位：__________",
        "若此單完整虧損，仍不會碰內部停止線與官方限制：□",
    ]
    for item in pretrade:
        add_text(doc, item, size=10, before=0, after=4, line=1.15)
    add_heading(doc, "交易後", 2)
    posttrade = [
        "實際結果：____ R／$__________　費用與滑價：__________",
        "完全依規則：□ 是　□ 否；偏差類型：__________",
        "是否需要策略變更：□ 否　□ 先記錄、完成樣本後評估",
        "截圖與決策文字已保存：□",
    ]
    for item in posttrade:
        add_text(doc, item, size=10, before=0, after=4, line=1.15)

    add_heading(doc, "附錄 B｜常見錯誤與修正", 1)
    mistakes = [
        ("用獲利目標決定倉位", "改用最大回撤與最差合理連敗決定風險。"),
        ("把 3% 當每日可用完的預算", "設定更低的內部停止線，保留滑價與錯單緩衝。"),
        ("只看 balance，不看 equity", "未實現虧損也會觸發 breach。"),
        ("用多個高度相關標的假裝分散", "按共同風險因子合併計算。"),
        ("回測好看就立刻購買", "先做樣本外、壓力測試與 forward test。"),
        ("連敗後加碼追回", "機械式日停損與降風險，不得翻倍。"),
        ("考核與 funded 使用不同策略", "同一策略、同一版本控制與紀錄。"),
        ("照抄網路策略或訊號", "只研究方法，交易假設與規則必須自行提出。"),
    ]
    add_table(doc, ["錯誤", "修正"], mistakes, [3900, 5460],
              [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    add_heading(doc, "來源與版本說明", 1)
    add_text(doc, "本手冊依 2026-07-27 可取得的官方資料整理。Breakout 可隨時更新規則、方案、價格、交易平台、標的或適用地區；購買前與每次交易前，均以你的合約、Program Rules 與 Dashboard 即時數字為準。")
    sources = [
        ("Breakout Program Rules（更新標示：2026-03-13）", "https://www.breakoutprop.com/program-rules/"),
        ("Breakout Evaluation Agreement", "https://www.breakoutprop.com/breakout-evaluation-agreement/"),
        ("Breakout Pricing／方案比較", "https://www.breakoutprop.com/pricing/"),
        ("How to Pass a Prop Firm Evaluation: The Sizing Math（2026-07-17）", "https://www.breakoutprop.com/article/pass-prop-firm-evaluation/"),
        ("Goodbye, DXtrade: What Changes on Breakout After July 9", "https://www.breakoutprop.com/article/dxtrade-retirement/"),
    ]
    for title, url in sources:
        p = doc.add_paragraph()
        set_para_format(p, 0, 2, 1.10)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "71")
        num_pr.append(ilvl)
        num_pr.append(num_id)
        p_pr.append(num_pr)
        add_hyperlink(p, title, url)
    # Metadata and core properties
    doc.core_properties.title = "Breakout Prop 制定交易策略的基礎知識"
    doc.core_properties.subject = "Breakout Evaluation 規則、風險管理、策略設計與驗證"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.keywords = "Breakout Prop, prop trading, crypto, risk management, strategy development"
    doc.core_properties.comments = "Rules checked against official sources on 2026-07-27."

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
